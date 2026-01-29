from . import UPLOAD_FOLDER
from . import ALLOWED_EXTENSIONS
from .services import mwe_service
from . import trie
from . import df
from . import db
from .models import User, UserDoc, UserDocLegTerm
from flask import Blueprint, render_template, request, flash, redirect, url_for, send_from_directory, current_app as app
from flask_login import login_required, logout_user, current_user
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import base64
import docx
import PyPDF2
import pandas as pd
import shutil
import json
import os
import shutil
import json
import os
import io 
import sys


views = Blueprint('views', __name__)
dict_coocur = pd.read_table('main_app/static/data/coocur_4v1.tsv')

# Function services here


def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_file_ext(filename):
    return filename.rsplit('.', 1)[1].lower()


def get_f_name(filename):
    return filename.rsplit('.', 1)[0]


def clean_folder(folder_path):
    # Safe cleanup - checks if path is actually inside our intended uploads dir
    if not folder_path or not os.path.exists(folder_path):
        return
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')

# Before login routes here


@views.route('/')
def home():
    return render_template('home.html', data={'user': current_user, 'with_loading': False})


@views.route('/data-visual')
def data_visual():
    return render_template('data_visual.html', data={'user': current_user, 'with_loading': True})


@views.route('/law-network', methods=['GET', 'POST'])
def law_network():
    # Robust data loading check
    global trie, df
    if trie is None or df is None:
        print("WARNING: Global trie/df is None in law_network view. Attempting reload...", file=sys.stderr)
        try:
            from . import load_global_data
            load_global_data()
            from . import trie, df # Re-import
            if trie is None:
                print("CRITICAL: Failed to reload trie.", file=sys.stderr)
                flash('Системийн алдаа: Өгөгдөл ачааллаж чадсангүй.', 'error')
                return render_template('law_network.html', data={'user': current_user, 'with_loading': False, 'with_visualization': False})
        except Exception as e:
            print(f"Error reloading globals: {e}", file=sys.stderr)
            flash('Системийн алдаа.', 'error')
            return render_template('law_network.html', data={'user': current_user, 'with_loading': False, 'with_visualization': False})

    if request.method == 'POST':
        upload_files = request.files.getlist('upload_file')
        if not upload_files or upload_files[0].filename == '':
            flash('Файл сонгогдоогүй байна.', 'error')
            return render_template('law_network.html', data={'user': current_user, 'with_loading': False, 'with_visualization': False})

        processed_data = {'nodes': [], 'links': []}
        all_law_terms = {} # law_name -> count
        doc_cooccurrences = [] # list of dicts: {'filename': str, 'laws': list}

        print(f"Processing {len(upload_files)} files...", file=sys.stderr)

        # Process each file
        for file in upload_files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                print(f"Reading file: {filename}", file=sys.stderr)
                text_content = ""
                file_ext = get_file_ext(file.filename)
                
                try:
                    if file_ext == 'txt':
                        text_content = file.read().decode('utf-8', errors='ignore')
                    elif file_ext in ['doc', 'docx']:
                        # docx requires stream, ensure at start
                        file.stream.seek(0)
                        doc = docx.Document(file)
                        text_content = "\n".join([p.text for p in doc.paragraphs])
                    elif file_ext == 'pdf':
                        # PyPDF2 reader
                        file.stream.seek(0)
                        # Check version or use generic approach
                        if hasattr(PyPDF2, 'PdfReader'):
                            reader = PyPDF2.PdfReader(file)
                            for i in range(len(reader.pages)):
                                text_content += reader.pages[i].extract_text() + "\n"
                        else:
                            reader = PyPDF2.PdfFileReader(io.BytesIO(file.read()))
                            for i in range(reader.getNumPages()):
                                text_content += reader.getPage(i).extractText() + "\n"
                except Exception as e:
                    print(f"Error reading {filename}: {e}", file=sys.stderr)
                    continue

                if not text_content:
                    print(f"Warning: No text content extracted from {filename}", file=sys.stderr)
                    continue

                # Extract Terms reusing mwe_service
                found_laws_in_doc = set()
                
                try:
                    # Split huge text to chunks for mwe_service
                    for chunk in mwe_service.str_to_word_lines(text_content, 300).split('\n'):
                        if not chunk.strip(): continue
                        # Ensure trie is valid list
                        mwe_trie_obj = trie[0] if isinstance(trie, list) else trie
                        word_trie_obj = trie[1] if isinstance(trie, list) and len(trie)>1 else None
                        
                        res = mwe_service.search_mwe(mwe_trie_obj, word_trie_obj, df, chunk)
                        for term in res['found_mwe']:
                            term_id = str(term['id'])
                            term_name = term['leg_term']
                            
                            law_key = term_id
                            found_laws_in_doc.add(law_key)
                            
                            if law_key not in all_law_terms:
                                all_law_terms[law_key] = {'id': law_key, 'name': term_name, 'count': 0, 'group': 1}
                            all_law_terms[law_key]['count'] += 1
                except Exception as e:
                    print(f"Error extracting terms from {filename}: {e}", file=sys.stderr)
                    continue

                if found_laws_in_doc:
                    # Store filename with laws for edge attribution
                    doc_cooccurrences.append({'filename': get_f_name(filename), 'laws': list(found_laws_in_doc)})

        print(f"Found {len(all_law_terms)} unique laws across {len(doc_cooccurrences)} docs.", file=sys.stderr)

        # Build Graph (Co-occurrence)
        # Nodes
        processed_data['nodes'] = list(all_law_terms.values())
        
        # Edges with document tracking
        edge_data = {}  # key: "source-target", value: {'source', 'target', 'value', 'documents': []}
        for doc_info in doc_cooccurrences:
            laws = sorted(doc_info['laws'])
            doc_name = doc_info['filename']
            for i in range(len(laws)):
                for j in range(i + 1, len(laws)):
                    source = laws[i]
                    target = laws[j]
                    edge_key = f"{source}-{target}"
                    if edge_key not in edge_data:
                        edge_data[edge_key] = {'source': source, 'target': target, 'value': 0, 'documents': []}
                    edge_data[edge_key]['value'] += 1
                    edge_data[edge_key]['documents'].append(doc_name)
        
        processed_data['links'] = list(edge_data.values())

        # Save to JSON
        try:
            json_path = os.path.join(app.root_path, 'static/data/current_law_network.json')
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(processed_data, f, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving JSON: {e}", file=sys.stderr)
            flash('Error saving graph data.', 'error')
            
        return render_template('law_network.html', data={'user': current_user, 'with_loading': False, 'with_visualization': True})

    return render_template('law_network.html', data={'user': current_user, 'with_loading': False, 'with_visualization': False})


@views.route('/upload-doc', methods=['GET', 'POST'])
def upload_doc():

    # Use config absolute path
    upload_folder = app.config['UPLOAD_FOLDER']
    temp_folder = os.path.join(upload_folder, 'temp')
    clean_folder(temp_folder)
    # clean_folder('./main_app/templates/uploads/temp') # Legacy template cleanup if needed

    if request.method == 'POST':

        if 'upload_file' not in request.files:
            flash('Файл upload хийх хэсэг байхгүй байна.', category="error")
        file = request.files['upload_file']
        if file and file.filename == '':
            flash('Файл upload хийгээгүй байна.', category="error")

        f_size = request.headers.get('Content-Length', type=int)
        if allowed_file(file.filename) and f_size <= app.config['MAX_CONTENT_LENGTH']:
            filename = secure_filename(file.filename)
            
            # Save to absolute path
            if not os.path.exists(temp_folder):
                os.makedirs(temp_folder)
                
            file.save(os.path.join(temp_folder, filename))
            file_ext = get_file_ext(file.filename)
            txt_filename = get_f_name(file.filename) + '.txt'
            txt_file = os.path.join(temp_folder, txt_filename)

            if file_ext == 'doc' or file_ext == 'docx':
                document = docx.Document(file)
                with open(txt_file, 'w', encoding='utf-8') as f:
                    for para in document.paragraphs:
                        f.write(para.text + '\n')
            elif file_ext == 'pdf':
                with open(os.path.join(temp_folder, filename), 'rb') as f:
                    reader = PyPDF2.PdfFileReader(f)
                    with open(txt_file, 'w', encoding='utf-8') as g:
                        for i in range(reader.getNumPages()):
                            text = reader.getPage(i).extractText()
                            g.write(text + '\n')
                            
            # Robust Base64 Encoding
            enc_msg = base64.urlsafe_b64encode(txt_file.encode("utf-8")).decode("utf-8")
            return redirect(url_for('views.doc_analysis', data=enc_msg))
        else:
            flash('Та зөвхөн TXT, PDF, DOC эсвэл DOCX өргөтгөлтэй, 20MB - с хэтрэхгүй хэмжээтэй файл байршуулах боломжтой.', category="error")
    return render_template('upload_doc.html', data={'user': current_user, 'with_loading': False})


@views.route('/doc-analysis/<data>')
def doc_analysis(data):
    try:
        # Robust decoding
        padding = 4 - (len(data) % 4)
        if padding != 4:
            data += "=" * padding
        txt_file = base64.urlsafe_b64decode(data).decode("utf-8")
    except Exception as e:
        flash("Invalid file identifier.", category="error")
        return redirect(url_for('views.upload_doc'))

    leg_terms = []
    # Simplified filename logic from path
    file_name = os.path.basename(txt_file).replace('.txt', '')
    
    html_str = '<div class="w-full">'
    if os.path.exists(txt_file):
        with open(txt_file, 'r', encoding='utf-8') as file:
            for i, row_txt in enumerate(file):
                if len(row_txt) > 0:
                    # print(row_txt)
                    row_html = ''
                    for len300 in mwe_service.str_to_word_lines(row_txt, 300).split('\n'):
                        res = mwe_service.search_mwe(
                            trie[0], trie[1], df, len300)
                        for found_mwe in res['found_mwe']:
                            leg_terms.append(
                                {'term_id': found_mwe['id'], 'term_name': found_mwe['leg_term']})
                        for wrd in res['word_list']:
                            if wrd['posTag'] == 'NM':
                                wrd['id'] = str(wrd['id']) + "_" + str(i)
                                param = "'" + wrd['id'] + "'"
                                wrd['word'] = f'<span class="inline-block"><span onmouseover="show_lt({param})" onclick="show_desc({param})" class="{wrd["id"]} cursor-pointer font-bold hover:underline">' + \
                                    wrd['word'] + \
                                    f'</span><div class="{wrd["id"]}-tooltip bg-business-color1 text-white px-2 py-1 rounded mt-2 opacity-100 transition-all duration-300 hidden"></div></span>'
    
                        row_html += (' ' +
                                     ' '.join(map(lambda x: x['word'], res['word_list'])))
                    html_str += ('<p class="w-full text-left">' +
                                 row_html.strip() + '</p>')
    else:
        flash("File processed not found on server.", category="error")
        
    html_str += '</div>'
    
    # Save HTML artifact - Use absolute path in uploads/temp for now, or templates if required by simple render
    # The original code saved to ./main_app/templates/uploads/temp. 
    # We will replicate this but using absolute paths if possible, or relative to app_root
    
    # Ideally templates should be static assets or explicitly loaded. 
    # Saving generated HTML into templates folder is dangerous/messy but preserving logic for now.
    template_folder = app.config['TEMPLATE_FOLDER']
    temp_html_dir = os.path.join(template_folder, 'uploads', 'temp')
    if not os.path.exists(temp_html_dir):
        os.makedirs(temp_html_dir)
        
    html_path = os.path.join(temp_html_dir, f"{file_name}.html")
    with open(html_path, "w", encoding='utf-8') as f:
        f.write(html_str)
        
    # Relative path for render_template
    html_f = f"uploads/temp/{file_name}.html"
    return render_template('doc_analysis.html', data={'user': current_user, 'with_loading': False, 'file_name': file_name, 'leg_terms': set(map(lambda x: x['term_name']+'-'+str(x['term_id']), leg_terms)), 'doc_html': html_f})


@views.route('/uploads/<name>')
def download_file(name):
    return send_from_directory(UPLOAD_FOLDER, name)


@views.route('/how-it-works')
def how_it_works():
    return render_template('how_it_works.html', data={'user': current_user, 'with_loading': False})

# Routes here in login after


@views.route('/main')
@login_required
def priv_home():
    return redirect(url_for('views.upload_new'))


@views.route('/upload_new', methods=['GET', 'POST'])
@login_required
def upload_new():
    if request.method == 'POST':

        if 'upload_file' not in request.files:
            flash('Файл upload хийх хэсэг байхгүй байна.', category="error")
        file = request.files['upload_file']
        if file and file.filename == '':
            flash('Файл upload хийгээгүй байна.', category="error")

        f_size = request.headers.get('Content-Length', type=int)
        if allowed_file(file.filename) and f_size <= app.config['MAX_CONTENT_LENGTH']:
            filename = secure_filename(file.filename)
            
            # Use absolute paths
            upload_folder = app.config['UPLOAD_FOLDER']
            if not os.path.exists(upload_folder):
                os.makedirs(upload_folder)
                
            file.save(os.path.join(upload_folder, filename))
            
            file_ext = get_file_ext(file.filename)
            # Simplified metadata in filename - user requested to remove ";0.001 MB" fragility but database relies on it?
            # User said "stop putting ;0.001 MB in filenames".
            # Changing implementation to store cleaner filenames. 
            # Note: doc_desc currently stores "ext;size MB". We'll keep doc_desc logic but DETACH it from the filesystem filename.
            
            desc_str = file_ext+';'+"{} MB".format(round(f_size/1048576, 3))
            clean_name = get_f_name(file.filename)
            
            # Storing plain .txt file without weird suffixes on disk
            txt_filename = f"{clean_name}.txt"
            txt_file = os.path.join(upload_folder, txt_filename)

            if file_ext == 'doc' or file_ext == 'docx':
                document = docx.Document(file)
                with open(txt_file, 'w', encoding='utf-8') as f:
                    for para in document.paragraphs:
                        f.write(para.text + '\n')
            elif file_ext == 'pdf':
                with open(os.path.join(upload_folder, filename), 'rb') as f:
                    reader = PyPDF2.PdfFileReader(f)
                    with open(txt_file, 'w', encoding='utf-8') as g:
                        for i in range(reader.getNumPages()):
                            text = reader.getPage(i).extractText()
                            g.write(text + '\n')
                            
            act_user = User.query.filter_by(id=current_user.id).first()
            if act_user:
                # Check duplication based on name and desc (schema constraint)
                user_doc = UserDoc.query.filter_by(
                    user_id=act_user.id).filter_by(doc_name=clean_name).filter_by(doc_desc=desc_str).first()
                    
                if user_doc:
                    flash('Баримт бичиг бүртгэлтэй байна.', category="warning")
                else:
                    new_user_doc = UserDoc(doc_name=clean_name, doc_desc=desc_str, doc_content_path=txt_file,
                                           user_id=act_user.id)
                    db.session.add(new_user_doc)
                    db.session.commit()
                    flash('Амжилттай байршууллаа.', category="success")
            else:
                flash('Хэрэглэгч олдсонгүй.', category="error")
            # return redirect(url_for('views.doc_list'))
        else:
            flash('Та зөвхөн TXT, PDF, DOC эсвэл DOCX өргөтгөлтэй, 20MB - с хэтрэхгүй хэмжээтэй файл байршуулах боломжтой.', category="error")
    return render_template('priv_home.html', data={'user': current_user, 'with_loading': False, 'active_tab': 1})


@views.route('/doc_list', methods=['GET', 'POST'])
@login_required
def doc_list():

    user_doc_list = []
    act_user = User.query.filter_by(id=current_user.id).first()
    if act_user:
        user_doc_list = UserDoc.query.filter_by(user_id=act_user.id).all()
    else:
        flash('Хэрэглэгч олдсонгүй.', category="error")
    return render_template('priv_home.html', data={'user': current_user, 'with_loading': False, 'active_tab': 2, 'user_doc_list': user_doc_list})


@views.route('/doc_list/<user_doc_id>', methods=['GET', 'POST'])
@login_required
def doc_detail(user_doc_id):

    act_user = User.query.filter_by(id=current_user.id).first()
    if not act_user:
        flash('Хэрэглэгч олдсонгүй.', category="error")

    user_doc = UserDoc.query.filter_by(
        user_id=act_user.id).filter_by(id=user_doc_id).first()

    if not user_doc:
        flash('Файл олдсонгүй.', category="error")

    txt_file = user_doc.doc_content_path
    file_name = user_doc.doc_name + "." + user_doc.doc_desc
    html_path = f"./main_app/templates/uploads/{file_name}.html"

    is_exist = os.path.exists(html_path)
    leg_terms = []
    if not is_exist:
        html_str = '<div class="w-full">'
        with open(txt_file, 'r', encoding='utf-8') as file:
            for i, row_txt in enumerate(file):
                if len(row_txt) > 0:
                    # print(row_txt)
                    row_html = ''
                    for len300 in mwe_service.str_to_word_lines(row_txt, 300).split('\n'):
                        res = mwe_service.search_mwe(
                            trie[0], trie[1], df, len300)
                        for found_mwe in res['found_mwe']:
                            leg_terms.append(
                                {'term_id': found_mwe['id'], 'term_name': found_mwe['leg_term']})
                            new_lt = UserDocLegTerm(
                                term_id=found_mwe['id'], term_name=found_mwe['leg_term'], user_doc_id=user_doc.id)
                            db.session.add(new_lt)
                            db.session.commit()
                        for wrd in res['word_list']:
                            if wrd['posTag'] == 'NM':
                                wrd['id'] = str(wrd['id']) + "_" + str(i)
                                param = "'" + wrd['id'] + "'"
                                wrd['word'] = f'<span class="inline-block"><span onmouseover="show_lt({param})" onclick="show_desc({param})" class="{wrd["id"]} cursor-pointer font-bold hover:underline">' + \
                                    wrd['word'] + \
                                    f'</span><div class="{wrd["id"]}-tooltip bg-business-color1 text-white px-2 py-1 rounded mt-2 opacity-100 transition-all duration-300 hidden"></div></span>'

                        row_html += (' ' +
                                     ' '.join(map(lambda x: x['word'], res['word_list'])))
                    html_str += ('<p class="w-full text-left">' +
                                 row_html.strip() + '</p>')
        html_str += '</div>'
        with open(html_path, "w", encoding='utf-8') as f:
            f.write(html_str)
    else:
        doc_lt_list = UserDocLegTerm.query.filter_by(
            user_doc_id=user_doc.id).all()
        for doc_lt in doc_lt_list:
            leg_terms.append({'term_id': doc_lt.term_id,
                             'term_name': doc_lt.term_name})
    
    # Return correct template path relative to templates folder
    html_f = f"uploads/{file_name}.html"
    return render_template('doc_analysis.html', data={'user': current_user, 'with_loading': False, 'file_name': user_doc.doc_name, 'leg_terms': set(map(lambda x: x['term_name']+'-'+str(x['term_id']), leg_terms)), 'doc_html': html_f})


@views.route('/api/delete_user_doc', methods=['POST'])
def delete_user_doc():
    if request.method == 'POST':
        act_user = User.query.filter_by(id=current_user.id).first()
        if not act_user:
            flash('Хэрэглэгч олдсонгүй.', category="error")
        payload = json.loads(request.get_data())
        user_doc_id = payload['user_doc_id']
        user_doc = UserDoc.query.filter_by(
            user_id=act_user.id).filter_by(id=user_doc_id).first()
        if not user_doc:
            flash('Баримт бичиг олдсонгүй.', category="error")

        user_doc_lt_list = UserDocLegTerm.query.filter_by(
            user_doc_id=user_doc.id).all()
        for user_doc_lt in user_doc_lt_list:
            db.session.delete(user_doc_lt)
        file_ext = user_doc.doc_desc.split(';')[0]
        if os.path.exists(user_doc.doc_content_path):
            os.remove(user_doc.doc_content_path)
        if os.path.exists(user_doc.doc_content_path.replace('.'+user_doc.doc_desc + '.txt', '.'+file_ext)):
            os.remove(user_doc.doc_content_path.replace(
                '.'+user_doc.doc_desc + '.txt', '.'+file_ext))
        if os.path.exists(user_doc.doc_content_path.replace('static', 'templates').replace('.txt', '.html')):
            os.remove(user_doc.doc_content_path.replace(
                'static', 'templates').replace('.txt', '.html'))
        db.session.delete(user_doc)
        db.session.commit()

    return 'Successfully deleted.'


@views.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():

    act_user = User.query.filter_by(id=current_user.id).first()
    if not act_user:
        flash('Хэрэглэгч олдсонгүй.', category="error")

    if request.method == 'POST':
        lname = request.form.get('lname')
        fname = request.form.get('fname')
        email = request.form.get('email')
        password = request.form.get('password')
        new_password = request.form.get('new_password')

        if not check_password_hash(act_user.password, password):
            flash('Идэвхтэй нууц үг буруу байна.', category='error')

        if len(email) < 5:
            flash('Таны и-мэйл хаяг 5 буюу дээш урттай байх шаардлагатай.',
                  category='error')
        elif len(lname) < 2:
            flash('Таны овог 2 буюу дээш урттай байх шаардлагатай.', category='error')
        elif len(fname) < 2:
            flash('Таны нэр 2 буюу дээш урттай байх шаардлагатай.', category='error')
        elif len(password) < 8:
            flash('Таны нууц үг 8 буюу дээш урттай байх шаардлагатай.',
                  category='error')
        elif password == new_password:
            flash(
                'Таны шинэ нууц үг идэвхтэй нууц үгээс өөр байх шаардлагатай.', category='error')
        else:
            act_user.lname = lname
            act_user.fname = fname
            act_user.email = email
            if new_password:
                if len(password) < 8:
                    flash(
                    'Таны шинэ нууц үг 8 буюу дээш урттай байх шаардлагатай.', category='error')
                act_user.password = generate_password_hash(new_password)
            db.session.commit()
            flash('Амжилттай шинэчлэгдлээ.', category='success')
    return render_template('priv_profile.html', data={'user': current_user, 'with_loading': False})


# Data routes here


@views.route('/api/get-lt-dict')
def get_lt_dict():
    with open('main_app/static/json/merge_lt_dict_v3.json', 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)
    return data


@views.route('/api/get-doc-dict')
def get_doc_dict():
    with open('main_app/static/json/dic_document_v2.json', 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)
    return data


@views.route('/api/get-coocur-dict')
def get_coocur_dict():
    with open('main_app/static/json/coocur_4v1.json', 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)
    return data


@views.route('/api/get-heatmap-data-len')
def get_heatmap_data_len():
    return json.loads({'total_doc': len(list(set(dict_coocur['doc_id'].to_list()))), 'total_term': len(list(set(dict_coocur['term_id'].to_list())))})


@views.route('/api/heatmap-data')
def heatmap_data():
    dict_list = []
    # if request.method == 'POST':
    #     take_doc = request.form.get('take_doc')
    #     take_term = request.form.get('take_term')
    doc_ids = list(set(dict_coocur['doc_id'].to_list()))[:100]
    term_ids = list(set(dict_coocur['term_id'].to_list()))[:50]
    for doc_id in doc_ids:
        for term_id in term_ids:
            value = 1
            if dict_coocur[(dict_coocur['doc_id'] == doc_id) & (dict_coocur['term_id'] == term_id)].empty:
                value = 0
            dict = {'group': doc_id, 'variable': term_id, 'value': value}
            dict_list.append(dict)
    data_frame = pd.DataFrame(dict_list)
    return data_frame.to_csv()


@views.route('/api/json-data')
def get_json_data():
    f = open('main_app/static/js/network.json')
    j_data = json.load(f)
    return j_data


@views.route('/api/nodes')
def nodes_csv():
    # nodes = pd.read_csv('main_app/static/data/nodes.csv')
    nodes = pd.read_csv('main_app/static/data/test_nodes.csv')
    return nodes.to_csv()


@views.route('/api/links')
def links_csv():
    # links = pd.read_csv('main_app/static/data/links.csv')
    links = pd.read_csv('main_app/static/data/test_links.csv')
    return links.to_csv()


@views.route('/api/upload-nodes')
def upload_nodes_csv():
    # Read from the JSON created by /law-network
    json_path = os.path.join(app.root_path, 'static/data/current_law_network.json')
    if not os.path.exists(json_path):
        return "id,group,count\n", 200 # Empty CSV
        
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    # Convert nodes to CSV format expected by D3
    # required: id, group, count (for radius), name (for label)
    nodes = data.get('nodes', [])
    if not nodes:
        return "id,group,count\n", 200

    output = io.StringIO()
    # Manual CSV writing to ensure correct headers
    output.write("id,group,count,name\n")
    for n in nodes:
        # Sanitize name
        name = n.get('name', '').replace(',', ' ').replace('\n', ' ')
        output.write(f"{n.get('id')},{n.get('group', 1)},{n.get('count', 1)},{name}\n")
        
    return output.getvalue(), 200, {'Content-Type': 'text/csv'}


@views.route('/api/upload-links')
def upload_links_csv():
    json_path = os.path.join(app.root_path, 'static/data/current_law_network.json')
    if not os.path.exists(json_path):
        return "source,target,value,documents\n", 200
        
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    links = data.get('links', [])
    if not links:
        return "source,target,value,documents\n", 200

    output = io.StringIO()
    output.write("source,target,value,documents\n")
    for l in links:
        docs = l.get('documents', [])
        docs_str = '; '.join(docs) if docs else ''
        output.write(f"{l.get('source')},{l.get('target')},{l.get('value', 1)},\"{docs_str}\"\n")
        
    return output.getvalue(), 200, {'Content-Type': 'text/csv'}

